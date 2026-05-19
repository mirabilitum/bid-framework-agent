# -*- coding: utf-8 -*-
"""
Document Generator
Generate Word document from framework JSON structure

Supports:
- [TABLE_START]...[TABLE_END] markers in content → rendered as Word tables
- 【xxx】 markers in content → rendered as bold section headers
"""

import json
import re
from typing import List, Dict, Any, Optional, Tuple


class DocumentGenerator:
    """Generate Word document from framework JSON"""

    FONT_NAME = '宋体'
    COVER_TITLE_SIZE = 18  # 小二 = 18pt
    TITLE_SIZE = 14  # 四号 = 14pt
    BODY_SIZE = 14

    def __init__(self, font_config: Optional[Dict[str, Any]] = None):
        """
        Args:
            font_config: Optional dict with keys:
                font_name (str), cover_title_size (int), title_size (int), body_size (int)
        """
        if font_config:
            self.FONT_NAME = font_config.get("font_name", self.FONT_NAME)
            self.COVER_TITLE_SIZE = font_config.get("cover_title_size", self.COVER_TITLE_SIZE)
            self.TITLE_SIZE = font_config.get("title_size", self.TITLE_SIZE)
            self.BODY_SIZE = font_config.get("body_size", self.BODY_SIZE)

    # Regex for 【xxx】 section headers in content
    _SECTION_HEADER_RE = re.compile(r'^【.+?】.*$')

    def generate(self, framework_nodes, output_path: str, project_name: str = ""):
        """
        Generate Word document from FrameworkNode list.

        Args:
            framework_nodes: list of FrameworkNode objects (or dict with "framework" key)
            output_path: output .docx path
            project_name: project name for main cover page
        """
        if isinstance(framework_nodes, dict):
            return self.generate_from_json(framework_nodes, output_path, project_name)

        # Convert FrameworkNode list to dict format
        def node_to_dict(node) -> dict:
            d = {
                "level": node.level,
                "title": node.title,
                "content": getattr(node, "content", ""),
                "children": [node_to_dict(c) for c in node.children],
            }
            if getattr(node, "cover_page", None):
                d["cover_page"] = node.cover_page
            if getattr(node, "index_page", None):
                d["index_page"] = node.index_page
            if getattr(node, "elements", None):
                d["elements"] = node.elements
            return d

        framework_data = {
            "has_chapter_covers": any(getattr(n, "cover_page", None) for n in framework_nodes),
            "framework": [node_to_dict(n) for n in framework_nodes],
        }
        return self.generate_from_json(framework_data, output_path, project_name)

    def generate_from_json(self, framework_data: Dict[str, Any], output_path: str, project_name: str = ""):
        """
        Generate Word document from framework JSON dict.

        Args:
            framework_data: dict with "has_chapter_covers" and "framework" keys
            output_path: output .docx path
            project_name: project name for main cover page
        """
        if "framework" not in framework_data or not framework_data["framework"]:
            raise ValueError("framework.json missing or empty 'framework' key")

        from docx import Document

        doc = Document()
        self._set_default_font(doc)

        has_covers = framework_data.get("has_chapter_covers", False)
        nodes = framework_data.get("framework", [])

        for node in nodes:
            has_cover = False
            if has_covers:
                cover = node.get("cover_page")
                if cover:
                    self._add_chapter_cover(doc, cover)
                    has_cover = True
                index = node.get("index_page")
                if index:
                    self._add_index_page(doc, index)

            if has_cover:
                # Cover page already shows section name — skip level 1 title,
                # render children directly
                for child in node.get("children", []):
                    child_level = child.get("level", 99)
                    self._add_node(doc, child, page_break=(child_level == 2))
            else:
                self._add_node(doc, node)

        doc.save(output_path)
        print(f"  [OK] Word document saved: {output_path}")

    def generate_from_file(self, json_path: str, output_path: str, project_name: str = ""):
        """Load framework JSON file and generate Word document."""
        with open(json_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        self.generate_from_json(data, output_path, project_name)

    # --- internal ---

    def _set_default_font(self, doc):
        """Set document default font to 宋体 四号"""
        from docx.shared import Pt
        style = doc.styles['Normal']
        font = style.font
        font.name = self.FONT_NAME
        font.size = Pt(self.BODY_SIZE)
        # For CJK font fallback
        style.element.get_or_add_rPr().rFonts.set(
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia',
            self.FONT_NAME
        )

    def _make_run(self, para, text: str, bold: bool = False, size: Optional[int] = None):
        """Add a run with 宋体 font to paragraph."""
        from docx.shared import Pt
        run = para.add_run(text)
        run.font.name = self.FONT_NAME
        run.font.size = Pt(size or self.BODY_SIZE)
        run.bold = bold
        run.element.get_or_add_rPr().rFonts.set(
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia',
            self.FONT_NAME
        )
        return run

    def _set_cell_font(self, cell, text: str, bold: bool = False):
        """Set font for a table cell."""
        from docx.shared import Pt
        # Clear default paragraph
        cell.text = ""
        p = cell.paragraphs[0]
        self._make_run(p, text, bold=bold)

    def _add_table(self, doc, header_line: str, data_lines: List[str], indent_level: int = 0):
        """
        Add a Word table from parsed [TABLE_START]...[TABLE_END] block.
        Supports merge syntax: [M:n] for column span, [B] for bold.
        Optional [COLS:n] as first line to declare total columns.

        Args:
            doc: Document object
            header_line: first line inside markers (column headers)
            data_lines: subsequent lines (data rows, may be empty)
            indent_level: node level for indentation context
        """
        import re
        from docx.shared import Pt, Inches, Cm
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn

        # Check for [COLS:n] declaration
        cols_match = re.match(r'\[COLS:(\d+)\]', header_line.strip())
        if cols_match:
            num_cols = int(cols_match.group(1))
            # header_line was [COLS:n], real header is first data_line
            if data_lines:
                header_line = data_lines[0]
                data_lines = data_lines[1:]
            else:
                return
        else:
            num_cols = None

        # Parse cells with merge/bold markers
        def parse_row(line):
            """Parse a row into list of (text, span, bold) tuples."""
            raw_cells = line.split('|')
            cells = []
            for raw in raw_cells:
                text = raw.strip()
                span = 1
                bold = False
                # Extract [M:n]
                m = re.search(r'\[M:(\d+)\]', text)
                if m:
                    span = int(m.group(1))
                    text = text[:m.start()] + text[m.end():]
                # Extract [B]
                if '[B]' in text:
                    bold = True
                    text = text.replace('[B]', '')
                cells.append((text.strip(), span, bold))
            return cells

        # Check if any line uses merge syntax
        all_lines = [header_line] + data_lines
        has_merge = any('[M:' in line or '[B]' in line for line in all_lines)

        if not has_merge and num_cols is None:
            # Simple table - preserve all cells including empty ones
            all_rows = [header_line] + data_lines
            parsed = []
            for line in all_rows:
                cells = [c.strip() for c in line.split('|')]
                parsed.append(cells)
            nc = max(len(row) for row in parsed)
            if nc == 0:
                return
            # Pad rows
            for row in parsed:
                while len(row) < nc:
                    row.append("")

            headers = parsed[0]
            rows_data = parsed[1:]

            table = doc.add_table(rows=len(parsed), cols=nc)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, h in enumerate(headers):
                self._set_cell_font(table.rows[0].cells[i], h, bold=True)
            for r_idx, row in enumerate(rows_data):
                for c_idx, val in enumerate(row):
                    self._set_cell_font(table.rows[r_idx + 1].cells[c_idx], val, bold=False)
            doc.add_paragraph()
            return

        # Merge table - parse all rows
        parsed_rows = [parse_row(line) for line in all_lines]

        # Determine num_cols from [COLS:n] or max row span sum
        if num_cols is None:
            num_cols = max(sum(span for _, span, _ in row) for row in parsed_rows)

        # Create table
        table = doc.add_table(rows=len(parsed_rows), cols=num_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Fill and merge
        for r_idx, row_cells in enumerate(parsed_rows):
            col_pos = 0
            for text, span, bold in row_cells:
                if col_pos >= num_cols:
                    break
                # Clamp span
                actual_span = min(span, num_cols - col_pos)
                cell = table.cell(r_idx, col_pos)
                if actual_span > 1:
                    merge_cell = table.cell(r_idx, col_pos + actual_span - 1)
                    cell.merge(merge_cell)
                self._set_cell_font(cell, text, bold=bold)
                col_pos += actual_span

        doc.add_paragraph()

    def _parse_content_blocks(self, content: str) -> List[Tuple[str, Any]]:
        """
        Parse content string into typed blocks:
        - ("text", line, align)           : normal text line with alignment
        - ("table", (header, [data_lines])) : table block
        - ("header", line)                : 【xxx】 section header (bold)

        align is one of: "left", "center", "right"

        Returns list of (type, data, ...) tuples.
        """
        blocks = []
        lines = content.split("\n")
        i = 0

        while i < len(lines):
            line = lines[i]

            # Check for [TABLE_START]
            if line.strip() == '[TABLE_START]':
                table_lines = []
                i += 1
                while i < len(lines) and lines[i].strip() != '[TABLE_END]':
                    table_lines.append(lines[i])
                    i += 1
                # i now points to [TABLE_END] or end
                if table_lines:
                    header = table_lines[0]
                    data = table_lines[1:]
                    blocks.append(("table", (header, data)))
                i += 1  # skip [TABLE_END]
                continue

            # Check for 【xxx】section header
            if self._SECTION_HEADER_RE.match(line.strip()):
                blocks.append(("header", line))
                i += 1
                continue

            # Check alignment markers: [CENTER] or [RIGHT]
            align = "left"
            text = line
            if line.startswith('[CENTER]'):
                align = "center"
                text = line[len('[CENTER]'):]
            elif line.startswith('[RIGHT]'):
                align = "right"
                text = line[len('[RIGHT]'):]

            blocks.append(("text", text, align))
            i += 1

        return blocks

    def _render_paragraphs(self, doc, paragraphs: List[Dict[str, Any]]):
        """Render paragraphs with precise formatting extracted from DOCX source."""
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        align_map = {
            "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
            "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
            "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
            "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        }

        for para_data in paragraphs:
            p = doc.add_paragraph()

            # Alignment
            align = para_data.get("align", "left")
            if align in align_map:
                p.alignment = align_map[align]

            # Indentation
            pf = p.paragraph_format
            left_indent = para_data.get("left_indent", 0)
            first_line = para_data.get("first_line_indent", 0)
            if left_indent:
                pf.left_indent = Cm(left_indent)
            if first_line:
                pf.first_line_indent = Cm(first_line)

            # Spacing
            space_before = para_data.get("space_before", 0)
            space_after = para_data.get("space_after", 0)
            if space_before:
                pf.space_before = Pt(space_before)
            if space_after:
                pf.space_after = Pt(space_after)

            # Text with bold and font size
            text = para_data.get("text", "")
            bold = para_data.get("bold", False)
            font_size = para_data.get("font_size")
            self._make_run(p, text, bold=bold, size=font_size)

    def _render_mixed_elements(self, doc, elements: List[Dict[str, Any]]):
        """Render mixed element list (paragraphs + tables) in document order."""
        for elem in elements:
            if elem.get("type") == "table":
                self._add_raw_table(doc, elem["rows"])
            else:
                # para — reuse paragraph rendering logic
                self._render_single_paragraph(doc, elem)

    def _render_single_paragraph(self, doc, para_data: Dict[str, Any]):
        """Render a single paragraph with precise formatting."""
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        align_map = {
            "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
            "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
            "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
            "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        }

        p = doc.add_paragraph()

        align = para_data.get("align", "left")
        if align in align_map:
            p.alignment = align_map[align]

        pf = p.paragraph_format
        left_indent = para_data.get("left_indent", 0)
        first_line = para_data.get("first_line_indent", 0)
        if left_indent:
            pf.left_indent = Cm(left_indent)
        if first_line:
            pf.first_line_indent = Cm(first_line)

        space_before = para_data.get("space_before", 0)
        space_after = para_data.get("space_after", 0)
        if space_before:
            pf.space_before = Pt(space_before)
        if space_after:
            pf.space_after = Pt(space_after)

        text = para_data.get("text", "")
        bold = para_data.get("bold", False)
        font_size = para_data.get("font_size")
        self._make_run(p, text, bold=bold, size=font_size)

    def _add_raw_table(self, doc, rows: List[List[str]]):
        """Render a Word table from raw row data extracted from DOCX source."""
        from docx.shared import Pt, Cm
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn

        if not rows or not rows[0]:
            return

        num_cols = max(len(row) for row in rows)
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Apply borders
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
        border_xml = (
            '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '</w:tblBorders>'
        )
        from lxml import etree
        tblPr.append(etree.fromstring(border_xml))

        for r_idx, row_data in enumerate(rows):
            row = table.rows[r_idx]
            for c_idx in range(num_cols):
                cell_text = row_data[c_idx] if c_idx < len(row_data) else ""
                is_header = (r_idx == 0)
                self._set_cell_font(row.cells[c_idx], cell_text, bold=is_header)

        # Add empty paragraph after table for spacing
        doc.add_paragraph()

    def _add_chapter_cover(self, doc, cover: Dict[str, Any]):
        """Add a chapter cover page (封面)."""
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        # Spacing before title
        for _ in range(6):
            doc.add_paragraph()

        # Title
        title_text = cover.get("title", "")
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self._make_run(p, title_text, bold=True, size=self.COVER_TITLE_SIZE)

        # Subtitle
        subtitle = cover.get("subtitle", "")
        if subtitle:
            p2 = doc.add_paragraph()
            p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            self._make_run(p2, subtitle, bold=True, size=16)

        # Spacing
        doc.add_paragraph()

        # Fields
        for field_text in cover.get("fields", []):
            p3 = doc.add_paragraph()
            p3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            self._make_run(p3, field_text)

        doc.add_page_break()

    def _add_index_page(self, doc, index: Dict[str, Any]):
        """Add an index page (索引)."""
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        # Index title
        title = index.get("title", "索引")
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self._make_run(p, title, bold=True)

        doc.add_paragraph()

        # Index items
        for item in index.get("items", []):
            p2 = doc.add_paragraph()
            self._make_run(p2, item)

        # Notes
        notes = index.get("notes", "")
        if notes:
            doc.add_paragraph()
            p3 = doc.add_paragraph()
            self._make_run(p3, notes)

        doc.add_page_break()

    def _add_node(self, doc, node: Dict[str, Any], page_break: bool = False):
        """Recursively add framework node. Alignment/indent from content markers only."""
        from docx.shared import Inches, Cm, Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        title = node.get("title", "")
        content = node.get("content", "")
        children = node.get("children", [])
        paragraphs = node.get("paragraphs")
        level = node.get("level", 99)

        # Page break before this node if requested
        if page_break:
            doc.add_page_break()

        # Title paragraph - use Heading styles for outline navigation
        title_align = None
        if title.startswith("[CENTER]"):
            title = title[len("[CENTER]"):]
            title_align = "center"

        # Map level to Heading style (1-3), deeper levels use bold paragraph
        heading_map = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3'}
        if level in heading_map and title:
            p_title = doc.add_paragraph(style=heading_map[level])
            if title_align == "center":
                p_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # Override heading style font to match document standard
            run = p_title.add_run(title)
            run.bold = True
            run.font.name = self.FONT_NAME
            run.font.size = Pt(self.TITLE_SIZE)
            run.font.color.rgb = RGBColor(0, 0, 0)
            from docx.oxml.ns import qn
            run._r.rPr.rFonts.set(qn('w:eastAsia'), self.FONT_NAME)
        else:
            p_title = doc.add_paragraph()
            if title_align == "center":
                p_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            self._make_run(p_title, title, bold=True)

        # Content rendering: prefer elements (DOCX-extracted mixed) > paragraphs > text markers
        elements = node.get("elements")
        if elements:
            self._render_mixed_elements(doc, elements)
        elif paragraphs:
            self._render_paragraphs(doc, paragraphs)
        elif content:
            blocks = self._parse_content_blocks(content)
            for block in blocks:
                block_type = block[0]

                if block_type == "table":
                    header_line, data_lines = block[1]
                    self._add_table(doc, header_line, data_lines, 0)

                elif block_type == "header":
                    # 【xxx】 rendered as bold line
                    p = doc.add_paragraph()
                    self._make_run(p, block[1], bold=True)

                else:
                    # ("text", text, align)
                    text = block[1]
                    align = block[2] if len(block) > 2 else "left"
                    p = doc.add_paragraph()
                    if align == "center":
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    elif align == "right":
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    # Leading spaces in text naturally create visual indent
                    self._make_run(p, text, bold=False)

        # Children — add page break before level 2 nodes for readability
        for child in children:
            child_level = child.get("level", 99)
            self._add_node(doc, child, page_break=(child_level == 2))
