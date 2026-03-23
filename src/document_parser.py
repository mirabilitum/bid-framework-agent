# -*- coding: utf-8 -*-
"""
Document Parser
Parse procurement documents (PDF/DOCX/DOC) and extract text, paragraphs, and tables.
Supports PDF page screenshots for vision-based format detection.
"""

import os
import re
import base64
from typing import Tuple, List, Dict, Optional


class DocumentParser:
    """Parse procurement documents (PDF / DOCX / DOC)."""

    def __init__(self, file_path: str):
        self.file_path = file_path
        self.ext = os.path.splitext(file_path)[1].lower()
        self.page_count: int = 0

    def parse(self) -> Tuple[str, List[dict], List[List[List[str]]]]:
        """
        Parse document.

        Returns:
            (full_text, paragraphs, tables)
            - full_text: concatenated document text
            - paragraphs: list of {"text": ..., "style": ...}
            - tables: list of tables, each table = list of rows, each row = list of cell strings
        """
        if self.ext == ".pdf":
            return self._parse_pdf()
        elif self.ext == ".docx":
            return self._parse_docx()
        elif self.ext == ".doc":
            return self._parse_doc()
        else:
            raise ValueError(f"Unsupported format: {self.ext}")

    # ------------------------------------------------------------------ #
    #  PDF page screenshots for vision-based format recognition
    # ------------------------------------------------------------------ #

    def screenshot_pages(self, page_numbers: List[int], dpi: int = 150) -> List[dict]:
        """
        Screenshot specified PDF pages as base64 PNG images.

        Args:
            page_numbers: 0-based page indices
            dpi: resolution (150 = good quality/token balance)

        Returns:
            [{"page": N, "data": base64_str, "media_type": "image/png"}, ...]
        """
        if self.ext != ".pdf":
            raise ValueError("screenshot_pages only supports PDF files")

        import fitz  # PyMuPDF

        doc = fitz.open(self.file_path)
        images = []
        for page_num in page_numbers:
            if page_num < 0 or page_num >= len(doc):
                continue
            page = doc[page_num]
            mat = fitz.Matrix(dpi / 72, dpi / 72)
            pix = page.get_pixmap(matrix=mat)
            png_data = pix.tobytes("png")
            images.append({
                "page": page_num,
                "data": base64.b64encode(png_data).decode("utf-8"),
                "media_type": "image/png",
            })
        doc.close()
        return images

    _CHAPTER_HEADING_RE = re.compile(r'第[一二三四五六七八九十\d]+章')

    def find_format_template_pages(self) -> List[int]:
        """
        Find page numbers containing format templates
        (e.g. chapter titled "投标文件格式" / "响应文件格式").
        Returns 0-based page indices.

        NOTE: This is a heuristic fallback. Prefer LLM-driven page detection
        (read full_text, identify format section, then use cli_screenshot.py).
        """
        if self.ext != ".pdf":
            return []

        import fitz

        doc = fitz.open(self.file_path)
        pages: List[int] = []
        in_section = False

        section_keywords = ["投标文件格式", "响应文件格式", "电子投标文件格式", "附件格式"]
        chapter_keywords = ["第七章", "第六章", "第八章", "附件"]

        for i in range(len(doc)):
            text = doc[i].get_text()

            if in_section:
                # Exit when hitting the next chapter
                if self._CHAPTER_HEADING_RE.search(text):
                    if not any(kw in text for kw in section_keywords):
                        break
                pages.append(i)
            else:
                if any(kw in text for kw in section_keywords):
                    if any(kw in text for kw in chapter_keywords):
                        in_section = True
                        pages.append(i)

        doc.close()
        return pages

    # ------------------------------------------------------------------ #
    #  DOCX format section: mixed element extraction (paragraphs + tables)
    # ------------------------------------------------------------------ #

    # Keywords that mark the start of a format template chapter
    _FORMAT_SECTION_KW = ["投标文件格式", "响应文件格式", "电子投标文件格式", "附件格式"]
    _CHAPTER_KW = ["第七章", "第六章", "第八章", "附件"]

    def extract_format_section_elements(self, start_keyword: Optional[str] = None) -> List[dict]:
        """
        按文档顺序提取格式章节所有元素（段落+表格）。

        定位逻辑：找到居中加粗且包含格式关键词的段落作为起点，
        遇到下一个"第X章"标题时停止。

        Returns:
            [
                {"type": "para", "text": "...", "align": "center", "bold": True, ...},
                {"type": "table", "rows": [["col1", "col2"], ["val1", "val2"]]},
                ...
            ]
        """
        if self.ext != ".docx":
            return []

        from docx import Document
        from docx.oxml.ns import qn

        doc = Document(self.file_path)
        body = doc.element.body

        # --- Locate section start paragraph ---
        # Must be centered + bold to avoid matching references in 投标须知
        start_elem = None
        for elem in body:
            if elem.tag.endswith('}p'):
                text = elem.text or ""
                # Collect full text from all runs
                full_text = "".join(
                    r.text or "" for r in elem.findall(qn('w:r'))
                )
                if not full_text.strip():
                    full_text = text
                full_text = full_text.strip()
                if not full_text:
                    continue

                # Check keyword match
                kw_match = False
                if start_keyword:
                    kw_match = start_keyword in full_text
                else:
                    kw_match = any(kw in full_text for kw in self._FORMAT_SECTION_KW)

                if not kw_match:
                    continue

                # Check centered alignment
                pPr = elem.find(qn('w:pPr'))
                jc = None
                if pPr is not None:
                    jc_elem = pPr.find(qn('w:jc'))
                    if jc_elem is not None:
                        jc = jc_elem.get(qn('w:val'))
                is_centered = jc == 'center'

                # Check bold
                is_bold = False
                for rPr in elem.findall(f'.//{qn("w:rPr")}'):
                    b_elem = rPr.find(qn('w:b'))
                    if b_elem is not None:
                        val = b_elem.get(qn('w:val'))
                        if val is None or val in ('1', 'true'):
                            is_bold = True
                            break

                if is_centered and is_bold:
                    start_elem = elem
                    break

        if start_elem is None:
            return []

        # --- Collect elements from start to next chapter ---
        result: List[dict] = []
        found_start = False

        # Build lookup dicts once to avoid O(n²) identity scans inside the loop
        para_map = {p._element: p for p in doc.paragraphs}
        table_map = {t._element: t for t in doc.tables}

        for elem in body:
            if elem is start_elem:
                found_start = True

            if not found_start:
                continue

            if elem.tag.endswith('}p'):
                para = para_map.get(elem)
                if para is None:
                    continue

                text = para.text.strip()

                # Stop at next chapter heading (but not the format section itself)
                if text and elem is not start_elem and self._CHAPTER_HEADING_RE.match(text):
                    if not any(kw in text for kw in self._FORMAT_SECTION_KW):
                        break

                fmt = self._extract_para_format(para)
                fmt["type"] = "para"
                result.append(fmt)

            elif elem.tag.endswith('}tbl'):
                tbl = table_map.get(elem)
                if tbl is not None:
                    rows = []
                    for row in tbl.rows:
                        rows.append([cell.text.strip() for cell in row.cells])
                    result.append({"type": "table", "rows": rows})

        return result

    # ------------------------------------------------------------------ #
    #  DOCX format template paragraph extraction (legacy, kept for compat)
    # ------------------------------------------------------------------ #

    def extract_format_section_paragraphs(
        self, start_keyword: Optional[str] = None
    ) -> List[dict]:
        """
        Extract paragraph-level formatting from a DOCX section as a flat list.

        Finds the section by start_keyword (or hardcoded fallback), then returns
        ALL paragraphs with their formatting until the next chapter heading.
        No template splitting — the LLM decides how to group them.

        Args:
            start_keyword: keyword to locate the section start (e.g. "投标文件格式").
                           If None, falls back to hardcoded _FORMAT_SECTION_KW.

        Returns:
            List of paragraph dicts (flat, in document order):
            [{"text", "align", "bold", "font_size", "left_indent",
              "first_line_indent", "space_before", "space_after"}, ...]
        """
        if self.ext != ".docx":
            return []

        from docx import Document
        from docx.shared import Emu

        doc = Document(self.file_path)
        paragraphs = doc.paragraphs

        # --- Locate section start ---
        start_idx = None
        for i, p in enumerate(paragraphs):
            text = p.text.strip()
            if start_keyword:
                if start_keyword in text:
                    start_idx = i
                    break
            else:
                if any(kw in text for kw in self._FORMAT_SECTION_KW):
                    if any(kw in text for kw in self._CHAPTER_KW):
                        start_idx = i
                        break

        if start_idx is None:
            return []

        # --- Extract all paragraphs until next chapter heading ---
        result: List[dict] = []
        # Include the heading paragraph itself
        result.append(self._extract_para_format(paragraphs[start_idx]))

        for p in paragraphs[start_idx + 1:]:
            text = p.text.strip()
            # Stop at next chapter heading (第X章)
            if text and self._CHAPTER_HEADING_RE.match(text):
                if not any(kw in text for kw in self._FORMAT_SECTION_KW):
                    break
            result.append(self._extract_para_format(p))

        return result

    def _extract_para_format(self, p) -> dict:
        """Extract formatting properties from a single DOCX paragraph."""
        from docx.shared import Emu

        # Alignment
        align_map = {0: "left", 1: "center", 2: "right", 3: "justify"}
        raw_align = p.alignment
        align = align_map.get(raw_align, "left") if raw_align is not None else "left"

        # Indentation (convert EMU to cm)
        pf = p.paragraph_format
        left_indent = round(pf.left_indent / Emu(360000), 2) if pf.left_indent else 0
        first_line = round(pf.first_line_indent / Emu(360000), 2) if pf.first_line_indent else 0

        # Spacing (convert EMU to pt: 1pt = 12700 EMU)
        space_before = round(pf.space_before / 12700, 1) if pf.space_before else 0
        space_after = round(pf.space_after / 12700, 1) if pf.space_after else 0

        # First run bold and font size
        bold = False
        font_size = None
        if p.runs:
            r0 = p.runs[0]
            bold = bool(r0.bold)
            if r0.font.size:
                font_size = round(r0.font.size / 12700, 1)  # EMU to pt

        return {
            "text": p.text,
            "align": align,
            "bold": bold,
            "font_size": font_size,
            "left_indent": left_indent,
            "first_line_indent": first_line,
            "space_before": space_before,
            "space_after": space_after,
        }

    # ------------------------------------------------------------------ #
    #  Format-specific parsers
    # ------------------------------------------------------------------ #

    def _parse_pdf(self) -> Tuple[str, List[dict], List[List[List[str]]]]:
        import fitz

        doc = fitz.open(self.file_path)
        self.page_count = len(doc)
        full_text = "\n".join(page.get_text() for page in doc)
        doc.close()

        if len(full_text.strip()) < 200:
            raise ValueError("PDF text too short — might be a scanned image")

        paragraphs = [{"text": full_text, "style": "Normal"}]

        # Extract tables using pdfplumber
        tables = []
        try:
            import pdfplumber
            with pdfplumber.open(self.file_path) as pdf:
                for page in pdf.pages:
                    for tbl in (page.extract_tables() or []):
                        cleaned = [[(cell or "").strip() for cell in row] for row in tbl]
                        if cleaned:
                            tables.append(cleaned)
        except ImportError:
            print("  [WARNING] pdfplumber not installed, skipping PDF table extraction")
        except Exception as e:
            print(f"  [WARNING] PDF table extraction failed: {e}")

        return full_text, paragraphs, tables

    def _parse_docx(self) -> Tuple[str, List[dict], List[List[List[str]]]]:
        from docx import Document

        doc = Document(self.file_path)

        paragraphs = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                paragraphs.append({
                    "text": text,
                    "style": p.style.name if p.style else "Normal",
                })

        full_text = "\n".join(p["text"] for p in paragraphs)

        tables = []
        for table in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            tables.append(rows)

        return full_text, paragraphs, tables

    def _parse_doc(self) -> Tuple[str, List[dict], List]:
        """Parse legacy .doc (requires Windows + MS Word COM)."""
        try:
            import win32com.client
            import pythoncom
        except ImportError:
            raise ImportError(
                "pywin32 is required for .doc parsing on Windows. "
                "Install with: pip install pywin32"
            )

        pythoncom.CoInitialize()
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(os.path.abspath(self.file_path))
            try:
                full_text = doc.Content.Text
            finally:
                doc.Close(False)
                word.Quit()
        finally:
            pythoncom.CoUninitialize()

        paragraphs = [{"text": full_text, "style": "Normal"}]
        return full_text, paragraphs, []
